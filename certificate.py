from docx import *
from docx.shared import *
import shutil
docPath = "C:\\Users\\Ibrahim.Moghul\\Desktop\\Data Analysis Scripts\\CSV OUTPUT\\TONGRUN\\Certificates\\"
docName = "test.docx"
xmlName = "test.xml"

def createCertificate(path,sn,cbDate,result):
    dest = path+"%s_certificate.docx"%sn
    shutil.copy2(path+"TEMPLATE.docx",dest)
    doc = Document(dest)

    for t in doc.tables:
        for c in t._cells:
            
            if c.text == "SERIAL NUMBER": 
                c.text = sn
                paragraphs = c.paragraphs
                paragraph = paragraphs[0]
                run_obj = paragraph.runs
                run = run_obj[0]
                font = run.font
                font.size = Pt(30)
                font.color.rgb = RGBColor(0xff,0xff,0xff)
                font.name = "AvenirNext LT Pro Regular"
            elif c.text == "CALIBRATION DATE" or c.text == "RESULT": 
                if(c.text == "CALIBRATION DATE"): c.text = cbDate
                elif (c.text == "RESULT"): c.text = result
                paragraphs = c.paragraphs
                paragraph = paragraphs[0]
                run_obj = paragraph.runs
                run = run_obj[0]
                font = run.font
                font.size = Pt(12)
                font.color.rgb = RGBColor(0xff,0xff,0xff)
                font.name = "AvenirNext LT Pro Regular"
    
            doc.save(dest)

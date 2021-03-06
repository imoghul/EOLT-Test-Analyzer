from docx import *
from docx.shared import *
import contextlib
from docx2pdf import convert
import os
import sys
import shutil


def docPath(sn, cbDate, path):
    return (
        path + sn + "_" + cbDate.replace("/", "") + "_certificate.docx"
    )  # path+"%s_%s_certificate.docx" % (sn,cbDate.replace("/",""))


def createCopy(sn, cbDate, path):
    dest = docPath(sn, cbDate, path)
    shutil.copy2(path + "TEMPLATE.docx", dest)


def createCertificate(sn, cbDate, result, DAQTemp, PostCalibAir, path, logger):
    try:
        dest = docPath(sn, cbDate, path)
        doc = Document(dest)

        for t in doc.tables:
            for c in t._cells:

                if c.text == "SERIAL NUMBER":
                    c.text = sn
                    paragraphs = c.paragraphs
                    paragraph = paragraphs[0]
                    run_obj = paragraph.runs
                    run = run_obj[0]
                    font = run.font
                    font.size = Pt(30)
                    font.color.rgb = RGBColor(0x00, 0x00, 0x00)
                    font.name = "AvenirNext LT Pro Regular"
                elif (
                    c.text == "CALIBRATION DATE"
                    or c.text == "RESULT"
                    or c.text == "DAQ TEMP"
                    or c.text == "CALIB"
                ):
                    try:
                        if c.text == "CALIBRATION DATE":
                            c.text = cbDate
                        elif c.text == "RESULT":
                            c.text = result
                        elif c.text == "DAQ TEMP":
                            c.text = str(float(DAQTemp))
                        elif c.text == "CALIB":
                            c.text = str(float(PostCalibAir))
                    except:
                        raise Exception(
                            "Error while generating the certificate")
                    paragraphs = c.paragraphs
                    paragraph = paragraphs[0]
                    run_obj = paragraph.runs
                    run = run_obj[0]
                    font = run.font
                    font.size = Pt(12)
                    font.color.rgb = RGBColor(0x00, 0x00, 0x00)
                    font.name = "AvenirNext LT Pro Regular"

                doc.save(dest)
    except Exception as e:
        # print(str(e))  # print("Couldn't generate certificate for "+sn)
        logger.error(e)


def convertToPDF_doc(doc):
    # with contextlib.redirect_stdout(open(os.devnull, 'w')):
    try:
        print("Converting %s to pdf" % doc.replace("\\", "/").split("/")[-1])
        convert(doc)
    except:
        print(doc + " couldn't be converted to a pdf")


def convertToPDF_path(path):
    # with contextlib.redirect_stdout(open(os.devnull, 'w')):
    try:
        print("Converting certificates to pdfs")
        convert(path)
    except Exception as e:
        print("Some files couldn't be converted to a pdf")
